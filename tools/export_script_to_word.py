#!/usr/bin/env python3
"""
ADVシーンのJSファイルからWord脚本を生成するツール

使い方:
  python3 tools/export_script_to_word.py tools/output/chapter7_code.js
  python3 tools/export_script_to_word.py tools/output/chapter7_code.js --out output/第七章脚本.docx
"""

import re
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx が必要です: pip3 install python-docx")
    sys.exit(1)


# ── カラー定義 ──────────────────────────────────────────
COLOR_HEADING   = RGBColor(0x1a, 0x3a, 0x5c)   # 濃紺（章・シーン見出し）
COLOR_SCENE_NUM = RGBColor(0x44, 0x44, 0x44)   # グレー（シーンID）
COLOR_SPEAKER_L = RGBColor(0x1a, 0x5c, 0x3a)   # 緑（左キャラ=ヤス）
COLOR_SPEAKER_R = RGBColor(0x5c, 0x1a, 0x1a)   # 赤茶（右キャラ）
COLOR_TEXT      = RGBColor(0x22, 0x22, 0x22)   # 本文


def parse_js_scenes(js_text: str) -> list[dict]:
    """
    JS ファイルから ADV シーンリストを抽出する。
    返り値: [{ 'id': 'c7s01', 'title': '...', 'script': [{'speaker':..., 'text':..., 'side':...}] }]
    """
    scenes = []

    # シーンブロックを抽出: cXsYY: { ... }  または  sceneXX: { ... }
    scene_pattern = re.compile(
        r'(?:^|\n)\s{0,4}(c\d+s\d+|scene\d+)\s*:\s*\{',
    )

    # シーン開始位置リストを取得
    positions = [(m.group(1), m.start()) for m in scene_pattern.finditer(js_text)]

    for i, (scene_id, start) in enumerate(positions):
        # このシーンブロックの終わりを次のシーン開始か EOF で区切る
        end = positions[i + 1][1] if i + 1 < len(positions) else len(js_text)
        block = js_text[start:end]

        # タイトル
        title_m = re.search(r"title\s*:\s*'([^']*)'", block)
        title = title_m.group(1) if title_m else ''

        # script 配列を抽出
        script_entries = []
        # speaker / text / side
        line_pattern = re.compile(
            r"\{\s*speaker\s*:\s*'([^']*)'\s*,\s*text\s*:\s*'([^']*)'\s*(?:,\s*side\s*:\s*'([^']*)')?"
        )
        for lm in line_pattern.finditer(block):
            script_entries.append({
                'speaker': lm.group(1),
                'text':    lm.group(2),
                'side':    lm.group(3) or '',
            })

        if script_entries:  # セリフがあるシーンのみ
            scenes.append({
                'id':     scene_id,
                'title':  title,
                'script': script_entries,
            })

    return scenes


def detect_chapter_title(js_text: str, js_path: Path) -> str:
    """JSファイル先頭コメントや定数から章タイトルを推測する。"""
    # // 第X章「...」 形式のコメント
    m = re.search(r'//\s*(第[一二三四五六七八九十\d]+章[「『]?[^」』\n]*[」』]?)', js_text[:300])
    if m:
        return m.group(1).strip()
    # ファイル名から
    stem = js_path.stem  # e.g. "chapter7_code"
    num_m = re.search(r'(\d+)', stem)
    if num_m:
        return f'第{num_m.group(1)}章'
    return js_path.stem


def set_font(run, size_pt: float, bold=False, color: RGBColor = None):
    run.font.name    = 'Hiragino Kaku Gothic ProN'
    run.font.size    = Pt(size_pt)
    run.font.bold    = bold
    if color:
        run.font.color.rgb = color


def build_word(scenes: list[dict], chapter_title: str, out_path: Path):
    doc = Document()

    # ── ページ設定 ──
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── スタイル調整 ──
    style = doc.styles['Normal']
    style.font.name = 'Hiragino Kaku Gothic ProN'
    style.font.size = Pt(10.5)

    # ── 章タイトル ──
    h = doc.add_heading('', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(chapter_title)
    set_font(run, 18, bold=True, color=COLOR_HEADING)
    doc.add_paragraph()  # 空行

    for scene in scenes:
        # ── シーン見出し ──
        p_scene = doc.add_paragraph()
        p_scene.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # シーンID
        id_run = p_scene.add_run(f'[{scene["id"]}]  ')
        set_font(id_run, 9, bold=False, color=COLOR_SCENE_NUM)
        # シーンタイトル
        title_run = p_scene.add_run(scene['title'] or '（タイトルなし）')
        set_font(title_run, 13, bold=True, color=COLOR_HEADING)
        # 下線
        p_scene.paragraph_format.space_before = Pt(14)
        p_scene.paragraph_format.space_after  = Pt(4)
        # 区切り線代わりに下線付きフォント（docxにはhrがないため）

        # ── セリフ ──
        for entry in scene['script']:
            speaker = entry['speaker']
            text    = entry['text']
            side    = entry['side']  # 'left' or 'right'

            p = doc.add_paragraph()
            p.paragraph_format.left_indent   = Cm(1.0)
            p.paragraph_format.space_before  = Pt(1)
            p.paragraph_format.space_after   = Pt(1)

            # 話者名
            spk_color = COLOR_SPEAKER_L if side == 'left' else COLOR_SPEAKER_R
            spk_run = p.add_run(f'{speaker}　')
            set_font(spk_run, 10, bold=True, color=spk_color)

            # セリフ本文
            txt_run = p.add_run(text)
            set_font(txt_run, 10.5, bold=False, color=COLOR_TEXT)

        doc.add_paragraph()  # シーン間の空行

    doc.save(str(out_path))
    print(f'✅ 保存しました: {out_path}')
    print(f'   シーン数: {len(scenes)}')
    total_lines = sum(len(s['script']) for s in scenes)
    print(f'   総セリフ行数: {total_lines}')


def main():
    parser = argparse.ArgumentParser(description='ADV脚本JSファイル → Wordドキュメント変換')
    parser.add_argument('js_file', help='入力JSファイル (例: tools/output/chapter7_code.js)')
    parser.add_argument('--out', '-o', help='出力Wordファイルパス (省略時は入力と同ディレクトリ)')
    args = parser.parse_args()

    js_path = Path(args.js_file)
    if not js_path.exists():
        print(f'エラー: ファイルが見つかりません: {js_path}')
        sys.exit(1)

    js_text = js_path.read_text(encoding='utf-8')

    # 出力パス決定
    if args.out:
        out_path = Path(args.out)
    else:
        # ファイル名から章番号を取得して img/ChapterX/ に保存
        # 例: chapter7_code.js → img/Chapter7/chapter7_脚本.docx
        num_m = re.search(r'(\d+)', js_path.stem)
        if num_m:
            chapter_num = num_m.group(1)
            # tools/output/ → tools/ → プロジェクトルート
            project_root = js_path.resolve().parent.parent.parent
            chapter_dir = project_root / 'img' / f'Chapter{chapter_num}'
            out_path = chapter_dir / f'chapter{chapter_num}_脚本.docx'
        else:
            out_path = js_path.parent / (js_path.stem.replace('_code', '') + '_脚本.docx')

    out_path.parent.mkdir(parents=True, exist_ok=True)

    chapter_title = detect_chapter_title(js_text, js_path)
    print(f'章タイトル: {chapter_title}')

    scenes = parse_js_scenes(js_text)
    print(f'パースしたシーン数: {len(scenes)}')

    if not scenes:
        print('シーンが見つかりませんでした。JSファイルの形式を確認してください。')
        sys.exit(1)

    build_word(scenes, chapter_title, out_path)


if __name__ == '__main__':
    main()
